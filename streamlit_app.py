import streamlit as st
import pandas as pd
from FunctionTools.version_one.common import common_structure
import traceback
import time
from datetime import datetime
from io import BytesIO
import markdown
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import html
from bs4 import BeautifulSoup

def clean_markdown_text(text):
    """Clean and prepare markdown text for conversion"""
    if not text:
        return ""
    
    # Remove markdown code blocks wrapper if present
    if text.strip().startswith('```markdown') and text.strip().endswith('```'):
        text = text.strip()[11:-3].strip()
    elif text.strip().startswith('```') and text.strip().endswith('```'):
        text = text.strip()[3:-3].strip()
    
    # Clean up excessive whitespace
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = text.strip()
    
    return text

def markdown_to_html(markdown_text):
    """Convert markdown to clean HTML"""
    if not markdown_text:
        return ""
    
    # Clean the markdown first
    cleaned_md = clean_markdown_text(markdown_text)
    
    # Convert markdown to HTML
    html_content = markdown.markdown(cleaned_md, extensions=['tables', 'fenced_code'])
    
    # Clean up the HTML with BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Remove empty paragraphs
    for p in soup.find_all('p'):
        if not p.get_text().strip():
            p.decompose()
    
    return str(soup)

def generate_pdf_report(all_results):
    """Generate a PDF report using ReportLab - pure Python, cloud-friendly"""
    
    try:
        st.info("🎨 Generating professional PDF with ReportLab...")
        
        # Create PDF buffer
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                               topMargin=0.8*inch, bottomMargin=0.8*inch,
                               leftMargin=0.8*inch, rightMargin=0.8*inch)
        
        # Get styles and create custom ones
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle('CustomTitle',
                                    parent=styles['Heading1'],
                                    fontSize=24,
                                    spaceAfter=30,
                                    textColor=colors.darkblue,
                                    alignment=TA_CENTER,
                                    fontName='Helvetica-Bold')
        
        heading_style = ParagraphStyle('CustomHeading',
                                      parent=styles['Heading2'],
                                      fontSize=16,
                                      spaceAfter=12,
                                      spaceBefore=20,
                                      textColor=colors.darkblue,
                                      fontName='Helvetica-Bold')
        
        subheading_style = ParagraphStyle('CustomSubheading',
                                         parent=styles['Heading3'],
                                         fontSize=14,
                                         spaceAfter=8,
                                         spaceBefore=15,
                                         textColor=colors.darkgreen,
                                         fontName='Helvetica-Bold')
        
        body_style = ParagraphStyle('CustomBody',
                                   parent=styles['Normal'],
                                   fontSize=11,
                                   spaceAfter=6,
                                   spaceBefore=3,
                                   alignment=TA_JUSTIFY,
                                   fontName='Helvetica')
        
        # Title
        story.append(Paragraph("Web Research Report", title_style))
        story.append(Spacer(1, 20))
        
        # Report metadata
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", body_style))
        story.append(Spacer(1, 10))
        
        # Summary table
        summary_data = [
            ['Report Summary', ''],
            ['Total Research Runs', str(len(all_results))],
            ['Successful Runs', str(len([r for r in all_results if 'error' not in r]))],
            ['Failed Runs', str(len([r for r in all_results if 'error' in r]))]
        ]
        
        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(summary_table)
        story.append(Spacer(1, 20))
        
        # Table of Contents
        story.append(Paragraph("Table of Contents", heading_style))
        story.append(Spacer(1, 10))
        
        for result_data in all_results:
            run_num = result_data["run_number"]
            company_name = result_data["company_name"]
            story.append(Paragraph(f"Research Run {run_num}: {company_name}", body_style))
        
        story.append(PageBreak())
        
        # Individual run sections
        for i, result_data in enumerate(all_results):
            if i > 0:
                story.append(PageBreak())
            
            run_num = result_data["run_number"]
            company_name = result_data["company_name"]
            country = result_data["country"]
            
            # Run header
            story.append(Paragraph(f"Research Run {run_num}: {company_name}", heading_style))
            story.append(Spacer(1, 15))
            
            # Company info table
            company_data = [
                ['Company Information', ''],
                ['Company Name', company_name],
                ['Country', country],
                ['Research Topic', result_data.get('research_topic', 'N/A')]
            ]
            
            if result_data.get("elapsed_minutes"):
                company_data.append(['Time Taken', f"{result_data['elapsed_minutes']:.2f} minutes"])
            
            company_table = Table(company_data, colWidths=[2*inch, 4*inch])
            company_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('BACKGROUND', (0, 1), (0, -1), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(company_table)
            story.append(Spacer(1, 15))
            
            # Search queries
            if result_data.get("search_queries"):
                story.append(Paragraph("Search Queries Used", subheading_style))
                for query in result_data["search_queries"]:
                    story.append(Paragraph(f"• {query}", body_style))
                story.append(Spacer(1, 10))
            
            # Support URLs
            if result_data.get("support_urls"):
                story.append(Paragraph("Support URLs Referenced", subheading_style))
                for url in result_data["support_urls"]:
                    story.append(Paragraph(f"• {url}", body_style))
                story.append(Spacer(1, 10))
            
            # Research results
            story.append(Paragraph("Research Results", subheading_style))
            
            if "error" in result_data:
                error_style = ParagraphStyle('ErrorStyle',
                                           parent=body_style,
                                           textColor=colors.red,
                                           fontName='Helvetica-Bold')
                story.append(Paragraph(f"Error: {result_data['error']}", error_style))
            else:
                if (result_data.get("result") and 
                    result_data["result"].get("final_data") and 
                    result_data["result"]["final_data"].get("web_response")):
                    
                    web_response = result_data["result"]["final_data"]["web_response"]
                    
                    # Process the markdown content
                    processed_content = process_markdown_for_reportlab(web_response, styles, story)
                    
                else:
                    story.append(Paragraph("No research results available for this run.", body_style))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        st.success("✅ Professional PDF generated successfully!")
        return buffer
        
    except Exception as e:
        st.error(f"❌ PDF Generation Error: {str(e)}")
        raise e

def process_markdown_for_reportlab(markdown_text, styles, story):
    """Process markdown content and add to ReportLab story"""
    
    # Clean the markdown
    cleaned_md = clean_markdown_text(markdown_text)
    
    # Convert to HTML first
    html_content = markdown.markdown(cleaned_md, extensions=['tables'])
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Custom styles for content
    content_heading_style = ParagraphStyle('ContentHeading',
                                          parent=styles['Heading3'],
                                          fontSize=13,
                                          spaceAfter=8,
                                          spaceBefore=12,
                                          textColor=colors.darkblue,
                                          fontName='Helvetica-Bold')
    
    content_body_style = ParagraphStyle('ContentBody',
                                       parent=styles['Normal'],
                                       fontSize=10,
                                       spaceAfter=6,
                                       spaceBefore=3,
                                       alignment=TA_JUSTIFY,
                                       fontName='Helvetica')
    
    # Process HTML elements
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'table', 'ul', 'ol']):
        try:
            if element.name in ['h1', 'h2', 'h3', 'h4']:
                text = element.get_text().strip()
                if text:
                    story.append(Paragraph(text, content_heading_style))
                    story.append(Spacer(1, 6))
                    
            elif element.name == 'p':
                text = element.get_text().strip()
                if text and len(text) > 2:  # Avoid empty paragraphs
                    # Handle special formatting
                    if text.startswith('**') and text.endswith('**'):
                        # Bold text
                        bold_style = ParagraphStyle('BoldText',
                                                   parent=content_body_style,
                                                   fontName='Helvetica-Bold')
                        story.append(Paragraph(text[2:-2], bold_style))
                    else:
                        story.append(Paragraph(text, content_body_style))
                    story.append(Spacer(1, 4))
                    
            elif element.name == 'table':
                # Convert HTML table to ReportLab table
                rows = element.find_all('tr')
                if rows:
                    table_data = []
                    for row in rows:
                        cells = row.find_all(['td', 'th'])
                        row_data = [cell.get_text().strip() for cell in cells]
                        if row_data:  # Only add non-empty rows
                            table_data.append(row_data)
                    
                    if table_data:
                        # Calculate column widths
                        num_cols = len(table_data[0]) if table_data else 1
                        col_width = 6.5 * inch / num_cols
                        
                        table = Table(table_data, colWidths=[col_width] * num_cols)
                        table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, -1), 9),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                            ('TOPPADDING', (0, 0), (-1, -1), 6),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black),
                            ('VALIGN', (0, 0), (-1, -1), 'TOP')
                        ]))
                        
                        story.append(table)
                        story.append(Spacer(1, 10))
                        
            elif element.name in ['ul', 'ol']:
                items = element.find_all('li')
                for item in items:
                    text = item.get_text().strip()
                    if text:
                        story.append(Paragraph(f"• {text}", content_body_style))
                        story.append(Spacer(1, 3))
                story.append(Spacer(1, 8))
                
        except Exception as element_error:
            # Skip problematic elements but continue processing
            continue
    
    return story

def generate_markdown_report(all_results):
    """Generate a markdown report that pandoc can convert beautifully"""
    
    # Start with title and metadata
    markdown_parts = [
        "---",
        "title: 'Web Research Report'",
        f"date: '{datetime.now().strftime('%B %d, %Y')}'",
        "author: 'Web Research Agent'",
        "geometry: 'margin=1in'",
        "fontsize: '11pt'",
        "documentclass: 'article'",
        "---",
        "",
        "\\newpage",
        "",
        "# Web Research Report",
        "",
        f"**Generated on:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## Report Summary",
        "",
        f"| Metric | Value |",
        f"|--------|-------|",
        f"| Total Research Runs | {len(all_results)} |",
        f"| Successful Runs | {len([r for r in all_results if 'error' not in r])} |",
        f"| Failed Runs | {len([r for r in all_results if 'error' in r])} |",
        "",
        "## Table of Contents",
        ""
    ]
    
    # Add table of contents
    for result_data in all_results:
        run_num = result_data["run_number"]
        company_name = result_data["company_name"]
        markdown_parts.append(f"- Research Run {run_num}: {company_name}")
    
    markdown_parts.extend(["", "\\newpage", ""])
    
    # Add individual run sections
    for i, result_data in enumerate(all_results):
        if i > 0:
            markdown_parts.append("\\newpage")
            markdown_parts.append("")
        
        run_num = result_data["run_number"]
        company_name = result_data["company_name"]
        country = result_data["country"]
        
        markdown_parts.extend([
            f"# Research Run {run_num}: {company_name}",
            "",
            "## Company Information",
            "",
            "| Field | Value |",
            "|-------|-------|",
            f"| Company Name | {company_name} |",
            f"| Country | {country} |",
            f"| Research Topic | {result_data.get('research_topic', 'N/A')} |"
        ])
        
        if result_data.get("elapsed_minutes"):
            markdown_parts.append(f"| Time Taken | {result_data['elapsed_minutes']:.2f} minutes |")
        
        markdown_parts.extend(["", ""])
        
        # Search queries
        if result_data.get("search_queries"):
            markdown_parts.extend([
                "## Search Queries Used",
                ""
            ])
            for query in result_data["search_queries"]:
                markdown_parts.append(f"- {query}")
            markdown_parts.append("")
        
        # Support URLs
        if result_data.get("support_urls"):
            markdown_parts.extend([
                "## Support URLs Referenced",
                ""
            ])
            for url in result_data["support_urls"]:
                markdown_parts.append(f"- {url}")
            markdown_parts.append("")
        
        # Research results
        markdown_parts.extend([
            "## Research Results",
            ""
        ])
        
        if "error" in result_data:
            markdown_parts.extend([
                f"**Error occurred:** {result_data['error']}",
                ""
            ])
        else:
            if (result_data.get("result") and 
                result_data["result"].get("final_data") and 
                result_data["result"]["final_data"].get("web_response")):
                
                web_response = result_data["result"]["final_data"]["web_response"]
                
                # Clean and format the markdown content
                cleaned_content = clean_markdown_text(web_response)
                
                markdown_parts.extend([
                    cleaned_content,
                    ""
                ])
            else:
                markdown_parts.extend([
                    "No research results available for this run.",
                    ""
                ])
    
    return "\n".join(markdown_parts)



def generate_html_report_for_pdf(all_results):
    """Generate HTML content optimized for xhtml2pdf conversion"""
    
    # CSS styles optimized for xhtml2pdf (inline styles work better)
    css_styles = """
    <style>
        @page {
            size: A4;
            margin: 2cm;
        }
        
        body {
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            font-size: 12pt;
        }
        
        .header {
            text-align: center;
            margin-bottom: 30px;
            border-bottom: 2px solid #2E86AB;
            padding-bottom: 15px;
        }
        
        .title {
            color: #2E86AB;
            font-size: 24pt;
            font-weight: bold;
            margin-bottom: 10px;
        }
        
        .subtitle {
            color: #666;
            font-size: 14pt;
            margin-bottom: 15px;
        }
        
        .metadata-table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }
        
        .metadata-table th, .metadata-table td {
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }
        
        .metadata-table th {
            background-color: #2E86AB;
            color: white;
            font-weight: bold;
        }
        
        .run-section {
            margin: 30px 0;
            page-break-inside: avoid;
            border: 1px solid #e0e0e0;
            padding: 15px;
            background-color: #fafafa;
        }
        
        .run-header {
            color: #2E86AB;
            font-size: 16pt;
            font-weight: bold;
            margin-bottom: 15px;
            border-bottom: 1px solid #2E86AB;
            padding-bottom: 5px;
        }
        
        .info-table {
            width: 100%;
            border-collapse: collapse;
            margin: 10px 0;
        }
        
        .info-table th, .info-table td {
            border: 1px solid #ddd;
            padding: 6px;
            text-align: left;
        }
        
        .info-table th {
            background-color: #f0f0f0;
            font-weight: bold;
            width: 30%;
        }
        
        .section-header {
            color: #1a5490;
            font-size: 14pt;
            font-weight: bold;
            margin: 20px 0 8px 0;
        }
        
        .content-section {
            margin: 15px 0;
            padding: 10px;
            background-color: white;
            border-left: 3px solid #2E86AB;
        }
        
        .query-list, .url-list {
            margin: 10px 0;
            padding-left: 20px;
        }
        
        .query-list li, .url-list li {
            margin: 5px 0;
            padding: 5px;
            background-color: #f8f9fa;
            border-radius: 3px;
        }
        
        .error-message {
            color: #dc3545;
            background-color: #f8d7da;
            border: 1px solid #f5c6cb;
            padding: 10px;
            margin: 10px 0;
            border-radius: 3px;
        }
        
        .research-content {
            line-height: 1.5;
        }
        
        .research-content h1, .research-content h2, .research-content h3 {
            color: #2E86AB;
            margin-top: 20px;
            margin-bottom: 10px;
        }
        
        .research-content table {
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }
        
        .research-content table th, .research-content table td {
            border: 1px solid #ddd;
            padding: 6px;
            text-align: left;
        }
        
        .research-content table th {
            background-color: #e3f2fd;
            font-weight: bold;
        }
        
        .page-break {
            page-break-before: always;
        }
    </style>
    """
    
    # Start building HTML
    html_parts = [
        "<!DOCTYPE html>",
        "<html>",
        "<head>",
        "<meta charset='UTF-8'>",
        "<title>Web Research Report</title>",
        css_styles,
        "</head>",
        "<body>",
        
        # Header section
        "<div class='header'>",
        "<div class='title'>Web Research Report</div>",
        f"<div class='subtitle'>Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>",
        "</div>",
        
        # Metadata table
        "<table class='metadata-table'>",
        "<tr><th>Report Information</th><th>Value</th></tr>",
        f"<tr><td>Total Research Runs</td><td>{len(all_results)}</td></tr>",
        f"<tr><td>Successful Runs</td><td>{len([r for r in all_results if 'error' not in r])}</td></tr>",
        f"<tr><td>Failed Runs</td><td>{len([r for r in all_results if 'error' in r])}</td></tr>",
        "</table>",
    ]
    
    # Table of Contents
    html_parts.extend([
        "<h2 class='section-header'>Table of Contents</h2>",
        "<ul>"
    ])
    
    for result_data in all_results:
        run_num = result_data["run_number"]
        company_name = result_data["company_name"]
        html_parts.append(f"<li>Research Run {run_num}: {html.escape(company_name)}</li>")
    
    html_parts.append("</ul>")
    
    # Individual run sections
    for i, result_data in enumerate(all_results):
        if i > 0:
            html_parts.append("<div class='page-break'></div>")
        
        run_num = result_data["run_number"]
        company_name = result_data["company_name"]
        country = result_data["country"]
        
        html_parts.extend([
            "<div class='run-section'>",
            f"<div class='run-header'>Research Run {run_num}: {html.escape(company_name)}</div>",
            
            # Basic information table
            "<table class='info-table'>",
            f"<tr><th>Company Name</th><td>{html.escape(company_name)}</td></tr>",
            f"<tr><th>Country</th><td>{html.escape(country)}</td></tr>",
            f"<tr><th>Research Topic</th><td>{html.escape(result_data.get('research_topic', 'N/A'))}</td></tr>",
        ])
        
        if result_data.get("elapsed_minutes"):
            html_parts.append(f"<tr><th>Time Taken</th><td>{result_data['elapsed_minutes']:.2f} minutes</td></tr>")
        
        html_parts.append("</table>")
        
        # Search queries
        if result_data.get("search_queries"):
            html_parts.extend([
                "<div class='section-header'>Search Queries Used</div>",
                "<ul class='query-list'>"
            ])
            for query in result_data["search_queries"]:
                html_parts.append(f"<li>{html.escape(query)}</li>")
            html_parts.append("</ul>")
        
        # Support URLs
        if result_data.get("support_urls"):
            html_parts.extend([
                "<div class='section-header'>Support URLs Referenced</div>",
                "<ul class='url-list'>"
            ])
            for url in result_data["support_urls"]:
                html_parts.append(f"<li>{html.escape(url)}</li>")
            html_parts.append("</ul>")
        
        # Research results
        html_parts.append("<div class='section-header'>Research Results</div>")
        
        if "error" in result_data:
            html_parts.append(f"<div class='error-message'>Error occurred: {html.escape(result_data['error'])}</div>")
        else:
            if (result_data.get("result") and 
                result_data["result"].get("final_data") and 
                result_data["result"]["final_data"].get("web_response")):
                
                web_response = result_data["result"]["final_data"]["web_response"]
                
                # Convert markdown content to HTML
                formatted_content = markdown_to_html(web_response)
                
                html_parts.extend([
                    "<div class='content-section'>",
                    "<div class='research-content'>",
                    formatted_content,
                    "</div>",
                    "</div>"
                ])
            else:
                html_parts.append("<p>No research results available for this run.</p>")
        
        html_parts.append("</div>")  # Close run-section
    
    # Close HTML
    html_parts.extend([
        "</body>",
        "</html>"
    ])
    
    return "\n".join(html_parts)

def generate_html_report(all_results):
    """Generate HTML content for the research report"""
    
    # CSS styles for better formatting
    css_styles = """
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 40px;
            color: #333;
        }
        
        .header {
            text-align: center;
            margin-bottom: 40px;
            border-bottom: 3px solid #2E86AB;
            padding-bottom: 20px;
        }
        
        .title {
            color: #2E86AB;
            font-size: 2.5em;
            margin-bottom: 10px;
            font-weight: bold;
        }
        
        .subtitle {
            color: #666;
            font-size: 1.2em;
            margin-bottom: 20px;
        }
        
        .metadata-table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            background: #f8f9fa;
        }
        
        .metadata-table th, .metadata-table td {
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }
        
        .metadata-table th {
            background: #2E86AB;
            color: white;
            font-weight: bold;
        }
        
        .run-section {
            margin: 40px 0;
            page-break-inside: avoid;
            border: 1px solid #e0e0e0;
            border-radius: 8px;
            padding: 20px;
            background: #fafafa;
        }
        
        .run-header {
            color: #2E86AB;
            font-size: 1.8em;
            margin-bottom: 20px;
            border-bottom: 2px solid #2E86AB;
            padding-bottom: 10px;
        }
        
        .info-table {
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }
        
        .info-table th, .info-table td {
            border: 1px solid #ddd;
            padding: 10px;
            text-align: left;
        }
        
        .info-table th {
            background: #f0f0f0;
            font-weight: bold;
            width: 30%;
        }
        
        .section-header {
            color: #1a5490;
            font-size: 1.3em;
            margin: 25px 0 10px 0;
            font-weight: bold;
        }
        
        .content-section {
            margin: 20px 0;
            padding: 15px;
            background: white;
            border-left: 4px solid #2E86AB;
            border-radius: 4px;
        }
        
        .query-list, .url-list {
            list-style-type: none;
            padding: 0;
        }
        
        .query-list li, .url-list li {
            margin: 8px 0;
            padding: 8px;
            background: #f8f9fa;
            border-radius: 4px;
            border-left: 3px solid #28a745;
        }
        
        .error-message {
            color: #dc3545;
            background: #f8d7da;
            border: 1px solid #f5c6cb;
            border-radius: 4px;
            padding: 15px;
            margin: 10px 0;
        }
        
        .research-content {
            line-height: 1.8;
            text-align: justify;
        }
        
        .research-content h1, .research-content h2, .research-content h3 {
            color: #2E86AB;
            margin-top: 25px;
            margin-bottom: 15px;
        }
        
        .research-content table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            background: white;
        }
        
        .research-content table th, .research-content table td {
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }
        
        .research-content table th {
            background: #e3f2fd;
            font-weight: bold;
        }
        
        .research-content table tr:nth-child(even) {
            background: #f8f9fa;
        }
        
        .page-break {
            page-break-before: always;
        }
    </style>
    """
    
    # Start building HTML
    html_parts = [
        "<!DOCTYPE html>",
        "<html>",
        "<head>",
        "<meta charset='UTF-8'>",
        "<title>Web Research Report</title>",
        css_styles,
        "</head>",
        "<body>",
        
        # Header section
        "<div class='header'>",
        "<h1 class='title'>Web Research Report</h1>",
        f"<p class='subtitle'>Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>",
        "</div>",
        
        # Metadata table
        "<table class='metadata-table'>",
        "<tr><th>Report Information</th><th>Value</th></tr>",
        f"<tr><td>Total Research Runs</td><td>{len(all_results)}</td></tr>",
        f"<tr><td>Successful Runs</td><td>{len([r for r in all_results if 'error' not in r])}</td></tr>",
        f"<tr><td>Failed Runs</td><td>{len([r for r in all_results if 'error' in r])}</td></tr>",
        "</table>",
    ]
    
    # Table of Contents
    html_parts.extend([
        "<h2 class='section-header'>Table of Contents</h2>",
        "<ul>"
    ])
    
    for result_data in all_results:
        run_num = result_data["run_number"]
        company_name = result_data["company_name"]
        html_parts.append(f"<li>Research Run {run_num}: {company_name}</li>")
    
    html_parts.append("</ul>")
    
    # Individual run sections
    for i, result_data in enumerate(all_results):
        if i > 0:
            html_parts.append("<div class='page-break'></div>")
        
        run_num = result_data["run_number"]
        company_name = result_data["company_name"]
        country = result_data["country"]
        
        html_parts.extend([
            "<div class='run-section'>",
            f"<h2 class='run-header'>Research Run {run_num}: {company_name}</h2>",
            
            # Basic information table
            "<table class='info-table'>",
            f"<tr><th>Company Name</th><td>{company_name}</td></tr>",
            f"<tr><th>Country</th><td>{country}</td></tr>",
            f"<tr><th>Research Topic</th><td>{result_data.get('research_topic', 'N/A')}</td></tr>",
        ])
        
        if result_data.get("elapsed_minutes"):
            html_parts.append(f"<tr><th>Time Taken</th><td>{result_data['elapsed_minutes']:.2f} minutes</td></tr>")
        
        html_parts.append("</table>")
        
        # Search queries
        if result_data.get("search_queries"):
            html_parts.extend([
                "<h3 class='section-header'>Search Queries Used</h3>",
                "<ul class='query-list'>"
            ])
            for query in result_data["search_queries"]:
                html_parts.append(f"<li>{html.escape(query)}</li>")
            html_parts.append("</ul>")
        
        # Support URLs
        if result_data.get("support_urls"):
            html_parts.extend([
                "<h3 class='section-header'>Support URLs Referenced</h3>",
                "<ul class='url-list'>"
            ])
            for url in result_data["support_urls"]:
                html_parts.append(f"<li><a href='{url}'>{html.escape(url)}</a></li>")
            html_parts.append("</ul>")
        
        # Research results
        html_parts.append("<h3 class='section-header'>Research Results</h3>")
        
        if "error" in result_data:
            html_parts.append(f"<div class='error-message'>Error occurred: {html.escape(result_data['error'])}</div>")
        else:
            if (result_data.get("result") and 
                result_data["result"].get("final_data") and 
                result_data["result"]["final_data"].get("web_response")):
                
                web_response = result_data["result"]["final_data"]["web_response"]
                
                # Convert markdown content to HTML
                formatted_content = markdown_to_html(web_response)
                
                html_parts.extend([
                    "<div class='content-section'>",
                    "<div class='research-content'>",
                    formatted_content,
                    "</div>",
                    "</div>"
                ])
            else:
                html_parts.append("<p>No research results available for this run.</p>")
        
        html_parts.append("</div>")  # Close run-section
    
    # Close HTML
    html_parts.extend([
        "</body>",
        "</html>"
    ])
    
    return "\n".join(html_parts)

def generate_docx_report(all_results):
    """Generate a DOCX report from research results with better markdown handling"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Web Research Report', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Report metadata
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total Research Runs: {len(all_results)}")
    doc.add_paragraph(f"Successful Runs: {len([r for r in all_results if 'error' not in r])}")
    doc.add_paragraph(f"Failed Runs: {len([r for r in all_results if 'error' in r])}")
    doc.add_paragraph()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    for result_data in all_results:
        run_num = result_data["run_number"]
        company_name = result_data["company_name"]
        doc.add_paragraph(f"Research Run {run_num}: {company_name}", style='List Number')
    
    doc.add_page_break()
    
    # Results for each run
    for i, result_data in enumerate(all_results):
        if i > 0:
            doc.add_page_break()
            
        run_num = result_data["run_number"]
        company_name = result_data["company_name"]
        country = result_data["country"]
        
        # Run header
        doc.add_heading(f'Research Run {run_num}: {company_name}', level=1)
        
        # Basic info table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        table.cell(0, 0).text = 'Company Name'
        table.cell(0, 1).text = company_name
        table.cell(1, 0).text = 'Country'
        table.cell(1, 1).text = country
        table.cell(2, 0).text = 'Research Topic'
        table.cell(2, 1).text = result_data.get('research_topic', 'N/A')
        
        if result_data.get("elapsed_minutes"):
            row = table.add_row()
            row.cells[0].text = 'Time Taken'
            row.cells[1].text = f"{result_data['elapsed_minutes']:.2f} minutes"
        
        doc.add_paragraph()
        
        # Search queries
        if result_data.get("search_queries"):
            doc.add_heading('Search Queries Used', level=2)
            for query in result_data["search_queries"]:
                doc.add_paragraph(query, style='List Bullet')
        
        # Support URLs
        if result_data.get("support_urls"):
            doc.add_heading('Support URLs Referenced', level=2)
            for url in result_data["support_urls"]:
                doc.add_paragraph(url, style='List Bullet')
        
        # Research results
        if "error" in result_data:
            doc.add_heading('Error', level=2)
            doc.add_paragraph(result_data['error'])
        else:
            doc.add_heading('Research Results', level=2)
            
            if (result_data.get("result") and 
                result_data["result"].get("final_data") and 
                result_data["result"]["final_data"].get("web_response")):
                
                web_response = result_data["result"]["final_data"]["web_response"]
                
                # Clean markdown and convert to proper format
                cleaned_content = clean_markdown_text(web_response)
                
                # Parse markdown content for DOCX
                html_content = markdown_to_html(cleaned_content)
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Process HTML elements and add to document
                for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'table', 'ul', 'ol']):
                    if element.name in ['h1', 'h2', 'h3', 'h4']:
                        level = int(element.name[1]) + 2  # Offset by 2 since we already have main headings
                        doc.add_heading(element.get_text().strip(), level=min(level, 9))
                    
                    elif element.name == 'p':
                        text = element.get_text().strip()
                        if text:
                            doc.add_paragraph(text)
                    
                    elif element.name == 'table':
                        # Convert HTML table to Word table
                        rows = element.find_all('tr')
                        if rows:
                            # Count maximum columns
                            max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
                            
                            # Create table
                            table = doc.add_table(rows=len(rows), cols=max_cols)
                            table.style = 'Table Grid'
                            
                            for i, row in enumerate(rows):
                                cells = row.find_all(['td', 'th'])
                                for j, cell in enumerate(cells):
                                    if j < max_cols:
                                        table.cell(i, j).text = cell.get_text().strip()
                    
                    elif element.name in ['ul', 'ol']:
                        list_items = element.find_all('li')
                        for item in list_items:
                            doc.add_paragraph(item.get_text().strip(), style='List Bullet' if element.name == 'ul' else 'List Number')
            else:
                doc.add_paragraph("No research results available for this run.")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit app configuration
st.set_page_config(page_title="Web Research Agent", layout="centered")

st.title("🔍 Web Researcher")
st.markdown("Enter the number of research runs and provide different parameters for each run.")

# Initialize session state for storing all states
if 'all_states' not in st.session_state:
    st.session_state.all_states = []

# Initialize session state for results persistence
if 'research_results' not in st.session_state:
    st.session_state.research_results = None

# Get number of research runs first
num_runs = st.number_input("Number of Research Runs", 
                           min_value=1, max_value=10, value=1, 
                           help="How many different research runs do you want to perform?")

# Adjust the all_states list size based on num_runs
if len(st.session_state.all_states) < num_runs:
    st.session_state.all_states.extend([None] * (num_runs - len(st.session_state.all_states)))
elif len(st.session_state.all_states) > num_runs:
    st.session_state.all_states = st.session_state.all_states[:num_runs]

# Create tabs or sections for each run's inputs
if num_runs > 0:
    st.markdown("---")
    st.subheader("📝 Input Parameters for Each Research Run")
    
    # Create input forms for each run
    for run_number in range(num_runs):
        with st.expander(f"🔍 Research Run {run_number + 1} Parameters", expanded=(num_runs <= 3)):
            with st.form(f"company_form_{run_number}"):
                st.markdown(f"**Run {run_number + 1} Configuration:**")
                
                company_name = st.text_input(f"Company Name", 
                                           placeholder="e.g., Infosys Limited", 
                                           key=f"company_{run_number}")
                country = st.text_input(f"Country", 
                                      placeholder="e.g., India", 
                                      key=f"country_{run_number}")
                research_topic = st.selectbox("Research Topic",
                                            options=["GENERAL", "NEWS", "FINANCE"],
                                            index=0,  # Default to "general" (first option)
                                            key=f"topic_{run_number}"
                                        )
                research_topic = research_topic.lower()
                search_queries = st.text_input(f"Search Queries", 
                                             placeholder="Enter your queries separated by forward slash, e.g., Question 1/Question 2", 
                                             key=f"queries_{run_number}")
                support_urls = st.text_input(f"Support URLs", 
                                           placeholder="Enter your links seperated by comma, e.g., https://www.xyz.com/home, https://www.xyz.com/contact", 
                                           key=f"urls_{run_number}")
                prompt = st.text_area(f"Prompt", 
                                    placeholder="Enter your prompt here", 
                                    height=150, 
                                    key=f"prompt_{run_number}")
                
                submitted = st.form_submit_button(f"Save Run {run_number + 1} Parameters")
                
                if submitted:
                    if company_name and country:
                        state = {
                            "company_name": company_name, 
                            "country": country, 
                            "research_topic": research_topic,
                            "search_queries": search_queries.split("/") if search_queries else None,
                            "support_urls": support_urls.split(",") if support_urls else None,
                            "prompt": prompt,
                            "run_number": run_number + 1
                        }
                        
                        # Save to session state
                        st.session_state.all_states[run_number] = state
                        
                        st.success(f"✅ Parameters saved for Run {run_number + 1}!")
                        print(f"Saved state for run {run_number + 1}:", state['company_name'])
                    else:
                        st.error("Please provide at least Company Name and Country.")

    # Display saved parameters summary
    if st.session_state.all_states and any(state is not None for state in st.session_state.all_states):
        st.markdown("---")
        st.subheader("📋 Saved Parameters Summary")
        
        saved_runs = [i+1 for i, state in enumerate(st.session_state.all_states) if state is not None]
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Runs Configured", len(saved_runs))
        with col2:
            st.metric("Total Runs Planned", num_runs)
        
        if saved_runs:
            st.write("**Configured Runs:**", ", ".join(map(str, saved_runs)))
            
            # Show brief summary of each saved run
            for i, state in enumerate(st.session_state.all_states):
                if state is not None:
                    st.write(f"**Run {i+1}:** {state['company_name']} ({state['country']})")

    # Execute all research runs
    if st.session_state.all_states and any(state is not None for state in st.session_state.all_states):
        st.markdown("---")
        
        if st.button("🚀 Execute All Research Runs", type="primary"):
            # Filter out None states
            valid_states = [state for state in st.session_state.all_states if state is not None]
            
            if not valid_states:
                st.error("No valid research configurations found. Please save parameters for at least one run.")
            else:
                # Initialize containers for results
                all_results = []
                
                # Progress bar for multiple runs
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                # Execute research for each saved state
                for idx, state in enumerate(valid_states):
                    run_number = state['run_number']
                    status_text.text(f"Running research {idx + 1} of {len(valid_states)} (Run {run_number})...")
                    
                    with st.spinner(f"Searching and extracting company information... (Run {run_number})"):
                        try:
                            print(f"Executing run {run_number} for:", state['company_name'])
                            start_time = time.time()
                            result = common_structure(company_name=state.get('company_name'),
                                                      country=state.get('country'),
                                                      research_topic=state.get('research_topic'),
                                                      search_queries=state.get('search_queries'),
                                                      prompt=state.get('prompt'),
                                                      support_urls=state.get('support_urls')
                                                      )
                            end_time = time.time()
                            elapsed_time = end_time - start_time
                            elapsed_minutes = elapsed_time / 60
                            
                            # Store result with run information
                            result_with_metadata = {
                                "run_number": run_number,
                                "company_name": state['company_name'],
                                "country": state['country'],
                                "research_topic": state['research_topic'],
                                "elapsed_minutes": elapsed_minutes,
                                "result": result,
                                "state": state,
                                "search_queries": state.get('search_queries', []),
                                "support_urls": state.get('support_urls', [])
                            }
                            all_results.append(result_with_metadata)
                            
                            # Update progress
                            progress_bar.progress((idx + 1) / len(valid_states))
                            
                        except Exception as e:
                            st.error(f"Error in run {run_number}: {str(e)}")
                            traceback.print_exc()
                            
                            # Store error result
                            error_result = {
                                "run_number": run_number,
                                "company_name": state['company_name'],
                                "country": state['country'],
                                "research_topic": state['research_topic'],
                                "elapsed_minutes": 0,
                                "error": str(e),
                                "result": None,
                                "state": state,
                                "search_queries": state.get('search_queries', []),
                                "support_urls": state.get('support_urls', [])
                            }
                            all_results.append(error_result)
                            
                            # Update progress even on error
                            progress_bar.progress((idx + 1) / len(valid_states))
                
                # Clear status text and progress bar
                status_text.empty()
                progress_bar.empty()
                
                # Store results in session state for persistence
                st.session_state.research_results = all_results
                
                # Display results
                if all_results:
                    st.success("✅ All research runs complete!")
                    
                    # Summary statistics
                    successful_runs = [r for r in all_results if "error" not in r]
                    failed_runs = [r for r in all_results if "error" in r]
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Total Runs", len(all_results))
                    with col2:
                        st.metric("Successful", len(successful_runs))
                    with col3:
                        st.metric("Failed", len(failed_runs))
                    
                    if successful_runs:
                        avg_time = sum(r["elapsed_minutes"] for r in successful_runs) / len(successful_runs)
                        st.metric("Average Time per Run", f"{avg_time:.2f} minutes")
                    
                    # Display results for each run
                    for i, result_data in enumerate(all_results):
                        run_num = result_data["run_number"]
                        company_name = result_data["company_name"]
                        country = result_data["country"]
                        
                        # Create expandable section for each run
                        with st.expander(f"📊 Run {run_num} Results: {company_name} ({country})", expanded=(len(all_results) == 1)):
                            if "error" in result_data:
                                st.error(f"❌ Run {run_num} failed: {result_data['error']}")
                            else:
                                result = result_data["result"]
                                elapsed_minutes = result_data["elapsed_minutes"]
                                
                                st.write(f"🕒 Time taken: {elapsed_minutes:.2f} minutes")
                                
                                structured_data = result.get("structured_data", {})
                                final_data = result.get("final_data", {})
                                
                                # Display the research results
                                st.subheader(f"{company_name}'s Research Results")
                                if final_data and "web_response" in final_data:
                                    st.write(final_data["web_response"])
                                else:
                                    st.write("No web response data available for this run.")
                                
                                # Optionally display structured data if available
                                if structured_data:
                                    st.subheader("📋 Structured Data")
                                    st.json(structured_data)
                else:
                    st.error("No results to display.")
    
# Download section - Show if we have results in session state
if st.session_state.research_results:
    st.markdown("---")
    st.subheader("📥 Download Research Report")
    
    all_results = st.session_state.research_results
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        download_format = st.selectbox(
            "Select Download Format:",
            options=["PDF", "DOCX"],
            index=0,
            help="Choose the format for your research report"
        )
    
    with col2:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if download_format == "PDF":
            try:
                with st.spinner("Generating professional PDF report..."):
                    pdf_buffer = generate_pdf_report(all_results)
                
                st.download_button(
                    label="📄 Download PDF Report",
                    data=pdf_buffer.getvalue(),
                    file_name=f"research_report_{timestamp}.pdf",
                    mime="application/pdf",
                    key="pdf_download"
                )
            except Exception as e:
                st.error(f"❌ PDF Generation Failed: {str(e)}")
                st.info("💡 This is a pure Python PDF generator that works on all platforms including Streamlit Cloud.")
        
        elif download_format == "DOCX":
            try:
                with st.spinner("Generating DOCX report..."):
                    docx_buffer = generate_docx_report(all_results)
                
                st.download_button(
                    label="📝 Download DOCX Report",
                    data=docx_buffer.getvalue(),
                    file_name=f"research_report_{timestamp}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="docx_download"
                )
            except Exception as e:
                st.error(f"❌ Error generating DOCX: {str(e)}")
    
    # Summary of available results
    st.info(f"💡 Report ready for download! Contains {len(all_results)} research runs with results from your last execution.")

# Show input section info if no results available yet
if not st.session_state.research_results:
    if st.session_state.all_states and any(state is not None for state in st.session_state.all_states):
        st.info("👆 Please click 'Execute All Research Runs' to generate results, then download reports will be available.")
    else:
        st.info("👆 Please configure parameters for at least one research run above, then execute to generate downloadable reports.")

# Optional: Add a button to clear all saved states and results
if (st.session_state.all_states and any(state is not None for state in st.session_state.all_states)) or st.session_state.research_results:
    st.markdown("---")
    if st.button("🗑️ Clear All Saved Parameters and Results", type="secondary"):
        st.session_state.all_states = []
        st.session_state.research_results = None
        st.success("All saved parameters and results cleared!")
        st.rerun()